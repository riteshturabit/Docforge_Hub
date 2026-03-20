from fastapi import FastAPI
import uuid
import json
from docx import Document
from notion_client import Client
from fastapi.responses import FileResponse
import os

from dotenv import load_dotenv

load_dotenv()
import html
from fastapi import HTTPException

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER

from backend.llm import llm
from backend.models import CompanyContext, GenerateSectionRequest
from backend.database import get_connection

app=FastAPI()

# Testing Purpose
@app.get("/")
def home():
    return {"message":"DocForge API Running"}


# Get Departments
@app.get("/departments")
def get_departments():

    conn=get_connection()
    cursor=conn.cursor()

    cursor.execute("SELECT * FROM departments")
    data=cursor.fetchall()

    cursor.close()
    conn.close()

    return {"departments":data}

# Get Templates
@app.get("/templates/{department_id}")
def get_templates(department_id : int):

    conn=get_connection()
    cursor=conn.cursor()

    cursor.execute(
        "SELECT id,name FROM document_templates WHERE department_id=%s",
        (department_id,)
    )

    data=cursor.fetchall()

    cursor.close()
    conn.close()

    return {"templates":data}

# Get Sections
@app.get("/sections/{template_id}")
def get_sections(template_id: int):

     conn=get_connection()
     cursor=conn.cursor()

     cursor.execute(
         """SELECT section_title, section_order FROM template_sections WHERE template_id=%s ORDER BY section_order""",
         (template_id,)
     )
     data=cursor.fetchall()

     cursor.close()
     conn.close()

     return {"sections": data}

# Company Context
@app.post("/company-context")
def save_company_context(data: CompanyContext):

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        INSERT INTO company_context
        (company_name, company_location, company_size,
        company_stage, product_type, target_customers,
        company_mission, company_vision)
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s)
        RETURNING id
        """,
        (
            data.company_name,
            data.company_location,
            data.company_size,
            data.company_stage,
            data.product_type,
            data.target_customers,
            data.company_mission,
            data.company_vision
        )
    )

    company_id = cursor.fetchone()[0]

    conn.commit()
    cursor.close()
    conn.close()

    return {"company_id": company_id}


# Get Company Context by ID
@app.get("/company-context/{company_id}")
def get_company_context(company_id: int):

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT 
            id,
            company_name,
            company_location,
            company_size,
            company_stage,
            product_type,
            target_customers,
            company_mission,
            company_vision,
            created_at
        FROM company_context
        WHERE id=%s
        """,
        (company_id,)
    )

    result = cursor.fetchone()

    if not result:
        raise HTTPException(status_code=404, detail="Company context not found")

    cursor.close()
    conn.close()

    return {
        "id": result[0],
        "company_name": result[1],
        "company_location": result[2],
        "company_size": result[3],
        "company_stage": result[4],
        "product_type": result[5],
        "target_customers": result[6],
        "company_mission": result[7],
        "company_vision": result[8],
        "created_at": str(result[9])
    }

# Create Document
@app.post("/create-document")
def create_document(template_id: int, company_id: int):

    conn = get_connection()
    cursor = conn.cursor()

    # Get actual template name
    cursor.execute(
        "SELECT name FROM document_templates WHERE id=%s",
        (template_id,)
    )
    result = cursor.fetchone()
    if not result:
        raise HTTPException(status_code=404, detail="Template not found")
    
    template_name = result[0]
    document_id = str(uuid.uuid4())

    cursor.execute(
        """
        INSERT INTO documents
        (id, template_id, company_id, title)
        VALUES (%s, %s, %s, %s)
        """,
        (document_id, template_id, company_id, template_name)
    )

    conn.commit()
    cursor.close()
    conn.close()

    return {"document_id": document_id, "title": template_name}


# Generate Questions

@app.post("/generate_questions")
def generate_questions(template_id : int):

    conn=get_connection()
    cursor=conn.cursor()

    #Get Doc Name
    cursor.execute(
        "SELECT name FROM document_templates WHERE id=%s",
        (template_id,)
    )
    result = cursor.fetchone()
    if not result:
        raise HTTPException(status_code=404, detail="Template not found")

    template_name = result[0]

    #Get Doc Section
    cursor.execute(
        """
        SELECT section_title, section_order
        FROM template_sections
        WHERE template_id=%s
        ORDER BY section_order
        """,
        (template_id,)
    )

    sections = cursor.fetchall()

    #Create Prompt
    prompt = f"""
You are an enterprise SaaS documentation assistant.

Generate around 40-45 questions required to create the following document.

Document: {template_name}

Sections:
{sections}

👉 IMPORTANT:
- Generate 2–3 questions PER SECTION
- Map each question to its section
- Total questions should be 40–45

Return the output ONLY in JSON format like this:

{{
 "sections":[
 {{
 "section":"Overview",
 "questions":[
  "question 1",
  "question 2",
  "question 3"
 ]
 }},
 {{
  "section":"Purpose",
  "questions":[
   "question 1",
   "question 2"
  ]
 }}
 ]
}}

Do not include explanations.
Only return JSON.
"""

    response=llm.invoke(prompt)
    try:
        data = json.loads(response.content)
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="LLM returned invalid JSON. Try again.")

    # Deleted old questions
    cursor.execute(
        "DELETE FROM template_questions WHERE template_id = %s",
        (template_id,)
    )

    # Insert new questions
    for sec in data["sections"]:

        section_title = sec["section"]

        cursor.execute(
            """
            SELECT section_order FROM template_sections
            WHERE template_id=%s AND section_title=%s
            """,
            (template_id, section_title)
        )

        result = cursor.fetchone()

        if not result:
            continue   # ✅ NOW VALID

        section_order = result[0]

        for i, q in enumerate(sec["questions"], start=1):
            cursor.execute(
                """
                INSERT INTO template_questions
                (template_id, section_title, question, section_order, question_order)
                VALUES (%s,%s,%s,%s,%s)
                """,
                (template_id, section_title, q, section_order, i)
            )

    conn.commit()
    cursor.close()
    conn.close()

    return {"message": "Questions generated and stored"}


# Next Questions

@app.get("/next_questions", name="Get Next Question Unique")
def get_next_questions(document_id: str, section_order: int):

    conn = get_connection()
    cursor = conn.cursor()

    # get template_id
    cursor.execute(
        "SELECT template_id FROM documents WHERE id=%s",
        (document_id,)
    )
    result = cursor.fetchone()
    if not result:
        raise HTTPException(status_code=404, detail="Document not found")
    template_id = result[0]

    # get questions for that section
    cursor.execute(
        """
        SELECT question
        FROM template_questions
        WHERE template_id=%s AND section_order=%s
        ORDER BY question_order
        """,
        (template_id, section_order)
    )

    questions = [row[0] for row in cursor.fetchall()]

    # get section name
    cursor.execute(
        """
        SELECT section_title
        FROM template_sections
        WHERE template_id=%s AND section_order=%s
        """,
        (template_id, section_order)
    )
    section= cursor.fetchone()

    cursor.close()
    conn.close()

    return {
        "section": section[0] if section else"",
        "questions": questions
    }


# Generate Sections

@app.post("/generate_section")
def generate_section(data: GenerateSectionRequest):

    conn = get_connection()
    cursor = conn.cursor()

    document_id = data.document_id
    section_order = data.section_order
    answers = data.answers

    # 1️⃣ Get template_id
    cursor.execute(
        "SELECT template_id FROM documents WHERE id=%s",
        (data.document_id,)
    )
    template_id = cursor.fetchone()[0]

    # 2️⃣ Get section title
    cursor.execute(
        """
        SELECT section_title
        FROM template_sections
        WHERE template_id=%s AND section_order=%s
        """,
        (template_id, data.section_order)
    )

    section_title = cursor.fetchone()[0]

    # 3️⃣ Format answers
    answers_text = "\n".join(
        [f"{a.question}: {a.answer}" for a in data.answers]
    )

    # 4️⃣ LLM prompt
    prompt = f"""
You are an enterprise SaaS documentation assistant.

Your job is to convert user-provided answers into a professional document section.

Generate professional content for the following section.

Section: {section_title}

User Answers:
{answers_text}

Guidelines:
- Keep it professional
- Strictly use the user's answers
- Expand only for clarity and professionalism
- Keep Content Aligned strictly with user answers
- Do not add assumptions or new policies
"""

    response = llm.invoke(prompt)
    content = response.content or "No content generated"


    cursor.execute(
        "DELETE FROM document_sections WHERE document_id=%s AND section_order=%s",
        (data.document_id, data.section_order)
    )

    cursor.execute(
        """
        INSERT INTO document_sections
        (document_id, section_title, section_content, section_order, is_completed)
        VALUES (%s,%s,%s,%s,TRUE)
        """,
        (data.document_id, section_title, content, data.section_order)
    )
    conn.commit()
    cursor.close()
    conn.close()

    return {
        "section": section_title,
        "content": content
    }

# # Save Section

# @app.post("/save_section")
# def save_section(data: dict):

#     conn = get_connection()
#     cursor = conn.cursor()

#     document_id = data["document_id"]
#     section_title = data["section_title"]
#     section_order = data["section_order"]
#     content = data["content"]

#     # 🔥 DELETE OLD (avoid duplicates)
#     cursor.execute(
#         """
#         DELETE FROM document_sections
#         WHERE document_id=%s AND section_order=%s
#         """,
#         (document_id, section_order)
#     )

#     # 🔥 INSERT NEW
#     cursor.execute(
#         """
#         INSERT INTO document_sections
#         (document_id, section_title, section_content, section_order, is_completed)
#         VALUES (%s,%s,%s,%s,TRUE)
#         """,
#         (document_id, section_title, content, section_order)
#     )

#     conn.commit()
#     cursor.close()
#     conn.close()

#     return {"message": "Section saved successfully"}


# Progress Bar
@app.get("/progress/{document_id}")
def get_progress(document_id: str):

    conn = get_connection()
    cursor = conn.cursor()

    # 1️⃣ get template_id
    cursor.execute(
        "SELECT template_id FROM documents WHERE id=%s",
        (document_id,)
    )
    template_id = cursor.fetchone()[0]

    # 2️⃣ total sections
    cursor.execute(
        """
        SELECT COUNT(*)
        FROM template_sections
        WHERE template_id=%s
        """,
        (template_id,)
    )
    total_sections = cursor.fetchone()[0]

    # 3️⃣ completed sections
    cursor.execute(
        """
        SELECT COUNT(*)
        FROM document_sections
        WHERE document_id=%s AND is_completed=TRUE
        """,
        (document_id,)
    )
    completed_sections = cursor.fetchone()[0]

    # 4️⃣ progress %
    progress = 0
    if total_sections > 0:
        progress = (completed_sections / total_sections) * 100

    cursor.close()
    conn.close()

    return {
        "completed_sections": completed_sections,
        "total_sections": total_sections,
        "progress": round(progress, 2)
    }

# Submit Answers

# @app.post("/submit_answers")
# def submit_answers(data: SubmitAnswersRequest):

#     conn=get_connection()
#     cursor=conn.cursor()

#     for item in data.answers:
#         cursor.execute(
#             """
#             INSERT INTO question_answers (document_id, questions, answer)
#             VALUES (%s,%s,%s)
#             """,
#             (
#                 str(data.document_id),
#                 item.question,
#                 item.answer
#             )
#         )
#     conn.commit()
#     cursor.close()
#     conn.close()

#     return {"message":"Answers saved successfully"}


# Generate Document
@app.post("/generate_document")
def generate_document(document_id: str):

    conn = get_connection()
    cursor=conn.cursor()

    # Delete Old sections
    cursor.execute(
        "DELETE FROM document_sections WHERE document_id=%s",
        (document_id,)
    )

    # Get template id
    cursor.execute(
        " SELECT template_id, company_id FROM documents WHERE id=%s",
        (document_id,)
    )
    doc = cursor.fetchone()

    template_id = doc[0]
    company_id = doc[1]

    
    # Get sections
    cursor.execute(
        """
        SELECT section_title, section_order
        FROM template_sections
        WHERE template_id=%s
        ORDER BY section_order
        """,
        (template_id,)
    )
    sections = cursor.fetchall()

    # Get answers
    cursor.execute(
        """
        SELECT questions,answer FROM question_answers
        WHERE document_id=%s
        """,
        (document_id,)
    )
    answers = cursor.fetchall()

    answers_text = "\n".join([f"{q}: {a}" for q,a in answers])

    generated_sections = []

    for section_title, section_order in sections:

        prompt = f"""
You are an Enterprise SaaS Documentation assistant.

User Answers : 
{answers_text}

Generate Professional Content for this section:

Section : {section_title}
"""
        response = llm.invoke(prompt)

        content = response.content

        cursor.execute(
            """
            INSERT INTO document_sections
            (document_id, section_title, section_content, section_order)
            VALUES (%s,%s,%s,%s)
            """,
            (document_id, section_title, content, section_order)
        )

        generated_sections.append(section_title)

    conn.commit()
    cursor.close()
    conn.close()

    return {
        "message" : "Document is Generated",
        "sections_created" : len(generated_sections)
    }


# Full Document

@app.get("/document/{document_id}")
def get_document(document_id: str):

    conn = get_connection()
    cursor = conn.cursor()

    # Get full document metadata
    cursor.execute(
        """
        SELECT 
            dt.name,
            d.version,
            d.status,
            d.created_at,
            d.notion_page_id,
            dep.name AS department,
            dty.name AS document_type,
            cc.company_name
        FROM documents d
        JOIN document_templates dt ON d.template_id = dt.id
        JOIN departments dep ON dt.department_id = dep.id
        JOIN document_types dty ON dt.document_type_id = dty.id
        LEFT JOIN company_context cc ON d.company_id = cc.id
        WHERE d.id = %s
        """,
        (document_id,)
    )

    meta = cursor.fetchone()

    if not meta:
        raise HTTPException(status_code=404, detail="Document not found")

    # Get sections
    cursor.execute(
        """
        SELECT DISTINCT ON (section_order)
            section_title,
            section_content,
            section_order
        FROM document_sections
        WHERE document_id=%s
        ORDER BY section_order, id DESC
        """,
        (document_id,)
    )

    rows = cursor.fetchall()

    cursor.close()
    conn.close()

    return {
        "id": document_id,
        "title": meta[0],
        "version": meta[1],
        "status": meta[2],
        "created_at": str(meta[3]),
        "is_published": meta[4] is not None,
        "notion_page_id": meta[4],
        "department": meta[5],
        "document_type": meta[6],
        "company_name": meta[7],
        "sections": [
            {
                "title": r[0],
                "content": r[1],
                "order": r[2]
            } for r in rows
        ]
    }

# Enhance Section

@app.post("/enhance_section")
def enhance_section(data: dict):

    conn = get_connection()
    cursor = conn.cursor()

    document_id = data["document_id"]
    section_order = data.get("section_order")
    action = data.get("action")
    custom_instruction = data.get("custom_instruction", "")


    if not document_id:
        raise HTTPException(status_code=400, detail="document_id required")


    # ACTION MAP
    action_map = {
        "longer": "Make the content more detailed and comprehensive",
        "shorter": "Make the content shorter and to the point",
        "formal": "Make the tone more formal and professional",
        "concise": "Make the content concise without losing meaning",
        "examples": "Add relevant examples",
        "table": "Add structured table if applicable",
        "clarity": "Improve clarity and readability",
        "grammar": "Fix grammar and improve sentence structure"
    }

    instruction = action_map.get(action, "")

    #  Single Section
    if section_order is not None:
        cursor.execute(
            """
            SELECT section_title, section_content
            FROM document_sections
            WHERE document_id=%s AND section_order=%s
            """,
            (document_id, section_order)
        )
        section = cursor.fetchone()

        if not section:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Section not found")

        section_title, content = section

        prompt = f"""
You are an Expert AI document editor.

Section: {section_title}

Content:
{content}

Instruction:
{instruction}
{custom_instruction}

Rules:
- Keep meaning same
- Improve quality only
- Do not add fake policies
- Return improved content only.
"""

        response = llm.invoke(prompt)

        cursor.close()
        conn.close()

        return {
            "section": section_title,
            "enhanced_content": response.content
        }

    #  FULL DOCUMENT
    else:
        cursor.execute(
            """
            SELECT section_title, section_content
            FROM document_sections
            WHERE document_id=%s
            ORDER BY section_order
            """,
            (document_id,)
        )

        sections = cursor.fetchall()

        if not sections:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Document Empty")

        full_text = "\n\n".join(
            [f"{s[0]}:\n{s[1]}" for s in sections]
        )

        prompt = f"""
You are an expert document editor.
Improve the full document.

Instruction:
{instruction}
{custom_instruction}

Document:
{full_text}

Return only improved document.
"""

        response = llm.invoke(prompt)

        cursor.close()
        conn.close()

        return {
            "enhanced_document": response.content
        }
    
# Save Enhance Section

@app.post("/save_enhanced_section")
def save_enhanced_section(data: dict):

    conn = get_connection()
    cursor = conn.cursor()

    document_id = data.get("document_id")
    section_order = data.get("section_order")
    content = data.get("content")

    cursor.execute(
        """
        UPDATE document_sections
        SET section_content=%s
        WHERE document_id=%s AND section_order=%s
        """,
        (content, document_id, section_order)
        
    )

    conn.commit()
    cursor.close()
    conn.close()

    return {"message": "Section updated successfully"}


# Download PDF

@app.get("/download/pdf/{document_id}")
def download_pdf(document_id: str):

    conn=get_connection()
    cursor=conn.cursor()

    # Doc Title
    cursor.execute(
        """
        SELECT dt.name
        FROM documents d
        JOIN document_templates dt ON d.template_id = dt.id
        WHERE d.id=%s
        """,
        (document_id,)
    )

    result = cursor.fetchone()

    if not result:
        raise HTTPException(status_code=404, detail="Document not found")

    title = result[0]

    # Get Sections
    cursor.execute(
        """
        SELECT DISTINCT ON (section_order)
        section_title, section_content, section_order
        FROM document_sections
        WHERE document_id=%s
        ORDER BY section_order, id DESC
        """,
        (document_id,)
    )

    sections = cursor.fetchall()

    if not sections:
        raise HTTPException(status_code=404, detail="No content found")

    # File Name
    file_name = title.lower().replace(" ", "_") + ".pdf"
    file_path = f"/tmp/{file_name}"

    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()

    content = []

    # Title Style
    title_style = styles["Heading1"]
    title_style.alignment = TA_CENTER

    content.append(Paragraph(html.escape(title), title_style))
    content.append(Spacer(1, 20))

    # Loop Sections
    for row in sections:

        sec_title = row[0]
        text = row[1]

        safe_title = sec_title if sec_title else "Untitled Section"
        safe_text = text if text else "No content available"

        safe_text = safe_text.replace("###", "")

        # Remove duplicate heading
        lines = safe_text.split("\n")

        if lines and lines[0].strip().lower() == safe_title.lower():
            lines = lines[1:]  # remove duplicate heading

        safe_text = "\n".join(lines).strip()

        safe_text = html.escape(safe_text)
        # Add Section Title
        content.append(Paragraph(f"<b>{safe_title}</b>", styles["Heading2"]))
        content.append(Spacer(1, 10))

        # Split lines 
        for line in safe_text.split("\n"):
            if line.strip():
                content.append(Paragraph(line, styles["Normal"]))
                content.append(Spacer(1, 6))

        content.append(Spacer(1, 12))

    doc.build(content)

    cursor.close()
    conn.close()

    return FileResponse(
        path=file_path,
        filename=file_name,
        media_type='application/pdf'
    )

# Download Docs

@app.get("/download/docx/{document_id}")
def download_docx(document_id: str):

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT section_title, section_content FROM document_sections
        WHERE document_id = %s
        ORDER BY section_order
        """, (document_id,)
    )

    sections = cursor.fetchall()

    doc = Document()

    for row in sections:

        sec_title = row[0]
        text = row[1]

        safe_title = sec_title if sec_title else "Untitled Section"
        safe_text = text if text else "No content available"
        
        safe_text = safe_text.replace("###", "")

        lines = safe_text.split("\n")

        if lines and lines[0].strip().lower() == safe_title.lower():
            lines = lines[1:]

        safe_text = "\n".join(lines).strip()

        doc.add_heading(safe_title, level=1)
        doc.add_paragraph("")

        for line in safe_text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

    file_path = f"/tmp/{document_id}.docx"
    doc.save(file_path)

    return FileResponse(
        path=file_path,
        filename=f"{document_id}.docx",
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

load_dotenv()

NOTION_TOKEN = os.getenv("NOTION_TOKEN")
NOTION_DB_ID = os.getenv("NOTION_DB_ID")

notion = Client(auth=NOTION_TOKEN)

#  Chunk helper
def chunk_blocks(blocks, size=100):
    for i in range(0, len(blocks), size):
        yield blocks[i:i + size]

# Push To Notion
@app.post("/push_to_notion")
def push_to_notion(document_id: str):

    conn = get_connection()
    cursor = conn.cursor()

    # STEP 1: Get all required fields including type, industry, version, department
    cursor.execute(
        """
        SELECT 
            d.title,
            d.created_at,
            d.version,
            dt.industry,
            dty.name AS doc_type,
            dep.name AS department
        FROM documents d
        JOIN document_templates dt ON d.template_id = dt.id
        JOIN document_types dty ON dt.document_type_id = dty.id
        JOIN departments dep ON dt.department_id = dep.id
        WHERE d.id = %s
        """,
        (document_id,)
    )

    result = cursor.fetchone()

    if not result:
        raise HTTPException(status_code=404, detail="Document not found")

    title       = result[0]
    created_at  = result[1]
    version     = result[2] or "v1.0"
    industry    = result[3] or "SaaS"
    doc_type    = result[4]
    department  = result[5]
    created_by  = "DocForge"

    # STEP 2: Get sections
    cursor.execute(
        """
        SELECT section_title, section_content
        FROM document_sections
        WHERE document_id=%s
        ORDER BY section_order
        """,
        (document_id,)
    )

    sections = cursor.fetchall()

    if not sections:
        raise HTTPException(status_code=400, detail="No sections to publish")

    # STEP 3: Build Notion blocks
    children = []

    for sec_title, sec_content in sections:

        children.append({
            "object": "block",
            "type": "heading_2",
            "heading_2": {
                "rich_text": [
                    {"type": "text", "text": {"content": sec_title}}
                ]
            }
        })

        if sec_content:
            for line in sec_content.split("\n"):
                line = line.strip()
                if line:
                    # Notion paragraph has 2000 char limit per block
                    if len(line) > 1990:
                        line = line[:1990]
                    children.append({
                        "object": "block",
                        "type": "paragraph",
                        "paragraph": {
                            "rich_text": [
                                {"type": "text", "text": {"content": line}}
                            ]
                        }
                    })

    # STEP 4: Create Notion page with ALL properties
    response = notion.pages.create(
        parent={"database_id": NOTION_DB_ID},
        properties={
            "Name": {
                "title": [
                    {"type": "text", "text": {"content": str(title)}}
                ]
            },
            "Type": {
                "select": {
                    "name": str(doc_type)
                }
            },
            "Industry": {
                "select": {
                    "name": str(industry)
                }
            },
            "Version": {
                "rich_text": [
                    {"type": "text", "text": {"content": str(version)}}
                ]
            },
            "Tags": {
                "multi_select": [
                    {"name": str(department)}
                ]
            },
            "Created_By": {
                "rich_text": [
                    {"type": "text", "text": {"content": created_by}}
                ]
            },
            "Created_at": {
                "date": {
                    "start": str(created_at)
                }
            }
        }
    )

    page_id = response["id"]

    # STEP 5: Append content blocks in chunks
    for block_chunk in chunk_blocks(children):
        notion.blocks.children.append(
            block_id=page_id,
            children=block_chunk
        )

    # STEP 6: Save notion_page_id back to DB
    cursor.execute(
        "UPDATE documents SET notion_page_id=%s WHERE id=%s",
        (page_id, document_id)
    )

    conn.commit()
    cursor.close()
    conn.close()

    return {
        "message": "Published to Notion",
        "notion_page_id": page_id,
        "title": title,
        "type": doc_type,
        "industry": industry,
        "version": version,
        "department": department
    }

# Get All Documents (Library Browser)
@app.get("/documents")
def get_all_documents(department_id: int = None):

    conn = get_connection()
    cursor = conn.cursor()

    if department_id:
        cursor.execute(
            """
            SELECT 
                d.id,
                d.title,
                d.version,
                d.status,
                d.created_at,
                d.notion_page_id,
                dt.name AS template_name,
                dt.industry,
                dep.name AS department_name,
                dty.name AS document_type,
                cc.company_name
            FROM documents d
            JOIN document_templates dt ON d.template_id = dt.id
            JOIN departments dep ON dt.department_id = dep.id
            JOIN document_types dty ON dt.document_type_id = dty.id
            LEFT JOIN company_context cc ON d.company_id = cc.id
            WHERE dep.id = %s
            ORDER BY d.created_at DESC
            """,
            (department_id,)
        )
    else:
        cursor.execute(
            """
            SELECT 
                d.id,
                d.title,
                d.version,
                d.status,
                d.created_at,
                d.notion_page_id,
                dt.name AS template_name,
                dt.industry,
                dep.name AS department_name,
                dty.name AS document_type,
                cc.company_name
            FROM documents d
            JOIN document_templates dt ON d.template_id = dt.id
            JOIN departments dep ON dt.department_id = dep.id
            JOIN document_types dty ON dt.document_type_id = dty.id
            LEFT JOIN company_context cc ON d.company_id = cc.id
            ORDER BY d.created_at DESC
            """
        )

    rows = cursor.fetchall()

    cursor.close()
    conn.close()

    documents = []
    for row in rows:
        documents.append({
            "id": row[0],
            "title": row[1],
            "version": row[2],
            "status": row[3],
            "created_at": str(row[4]),
            "notion_page_id": row[5],
            "is_published": row[5] is not None,
            "template_name": row[6],
            "industry": row[7],
            "department": row[8],
            "document_type": row[9],
            "company_name": row[10]
        })

    return {
        "total": len(documents),
        "documents": documents
    }