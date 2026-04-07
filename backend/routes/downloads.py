import html
import re
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from backend.database import get_connection
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak, CondPageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.lib.units import inch
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

router = APIRouter()

# Shared helpers 

def render_rich_text_pdf(text: str) -> str:
    text = re.sub(
        r'^([A-Za-z][A-Za-z0-9\s\(\)\/\-\&]{2,120}):\s',
        r'**\1:** ',
        text
    )
    text = html.escape(text)
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.*?)__',     r'<u>\1</u>', text)
    return text


def clean_text_basic(text: str) -> str:
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'__(.*?)__',     r'\1', text)
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.*?)_{1,3}',   r'\1', text)
    text = re.sub(r'`{1,3}(.*?)`{1,3}',   r'\1', text)
    text = re.sub(r'^[-•]\s+', '', text)
    return text.strip()


def is_table_row(line: str) -> bool:
    line = line.strip()
    return line.startswith('|') and line.endswith('|')


def is_separator_row(line: str) -> bool:
    return bool(re.match(r'\|[-:\s|]+\|', line.strip()))


def parse_table_rows(lines: list) -> list:
    rows = []
    for line in lines:
        if is_separator_row(line):
            continue
        if is_table_row(line):
            cells = [clean_text_basic(c.strip()) for c in line.strip('|').split('|')]
            cells = [c for c in cells if c != '']
            if cells:
                rows.append(cells)
    return rows


# PDF DOWNLOAD

@router.get("/download/pdf/{document_id}")
def download_pdf(document_id: str):
    conn   = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT dt.name FROM documents d
        JOIN document_templates dt ON d.template_id = dt.id
        WHERE d.id=%s
        """,
        (document_id,)
    )
    result = cursor.fetchone()
    if not result:
        raise HTTPException(status_code=404, detail="Document not found")
    title = result[0]

    cursor.execute(
        """
        SELECT DISTINCT ON (section_order)
            section_title, section_content, section_order
        FROM document_sections
        WHERE document_id=%s
        ORDER BY section_order, id DESC
        """,
        (document_id,)
    )
    sections = cursor.fetchall()
    if not sections:
        raise HTTPException(status_code=404, detail="No content found")

    cursor.close()
    conn.close()

    # Get cover page metadata 
    conn2   = get_connection()
    cursor2 = conn2.cursor()
    cursor2.execute(
        """
        SELECT
            dep.name,
            dty.name,
            cc.company_name,
            dt.industry,
            d.version,
            d.created_at
        FROM documents d
        JOIN document_templates dt ON d.template_id = dt.id
        JOIN departments dep ON dt.department_id = dep.id
        JOIN document_types dty ON dt.document_type_id = dty.id
        LEFT JOIN company_context cc ON d.company_id = cc.id
        WHERE d.id = %s
        """,
        (document_id,)
    )
    meta_row = cursor2.fetchone()
    cursor2.close()
    conn2.close()

    department   = meta_row[0] if meta_row else ""
    doc_type     = meta_row[1] if meta_row else "Document"
    company_name = meta_row[2] if meta_row else ""
    industry     = meta_row[3] if meta_row else ""
    version      = meta_row[4] if meta_row else "v1.0"
    created_at   = meta_row[5] if meta_row else datetime.now()
    date_str     = created_at.strftime("%B %d, %Y") if created_at else ""

    file_name = title.lower().replace(" ", "_") + ".pdf"
    file_path = f"/tmp/{file_name}"

    doc   = SimpleDocTemplate(
        file_path,
        rightMargin=60, leftMargin=60,
        topMargin=60,   bottomMargin=80
    )
    styles = getSampleStyleSheet()
    story  = []

    # Content styles 
    section_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Heading2'],
        fontSize=14,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#7F77DD'),
        spaceBefore=16,
        spaceAfter=6,
        keepWithNext=True
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica',
        textColor=colors.HexColor('#2a2a4a'),
        leading=18,
        spaceAfter=6,
        alignment=TA_JUSTIFY
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica',
        textColor=colors.HexColor('#2a2a4a'),
        leading=18,
        spaceAfter=4,
        leftIndent=16,
        firstLineIndent=0,
        alignment=TA_JUSTIFY
    )
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontSize=10,
        fontName='Helvetica-Bold',
        textColor=colors.white
    )
    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontSize=10,
        fontName='Helvetica',
        textColor=colors.HexColor('#2a2a4a')
    )
    cover_doctype_style = ParagraphStyle(
        'CoverDocType',
        parent=styles['Normal'],
        fontSize=10,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#7F77DD'),
        alignment=TA_CENTER,
        spaceAfter=12,
        leading=14
    )
    cover_title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Normal'],
        fontSize=28,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a1a3a'),
        alignment=TA_CENTER,
        spaceAfter=8,
        leading=36
    )
    cover_company_style = ParagraphStyle(
        'CoverCompany',
        parent=styles['Normal'],
        fontSize=15,
        fontName='Helvetica',
        textColor=colors.HexColor('#4a4a6a'),
        alignment=TA_CENTER,
        spaceAfter=6,
        leading=20
    )
    cover_desc_style = ParagraphStyle(
        'CoverDesc',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica',
        textColor=colors.HexColor('#888888'),
        alignment=TA_CENTER,
        spaceAfter=6,
        leading=17
    )
    meta_label_style = ParagraphStyle(
        'MetaLabel',
        parent=styles['Normal'],
        fontSize=8,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#7F77DD'),
        leading=12
    )
    meta_value_style = ParagraphStyle(
        'MetaValue',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1a1a3a'),
        leading=15
    )

    # COVER PAGE

    story.append(HRFlowable(
        width="100%", thickness=8,
        color=colors.HexColor('#7F77DD'),
        spaceAfter=40
    ))

    logo_inner = Table(
        [[Paragraph('<b>D</b>', ParagraphStyle(
            'LogoLetter',
            parent=styles['Normal'],
            fontSize=16,
            fontName='Helvetica-Bold',
            textColor=colors.white,
            alignment=TA_CENTER
        ))]],
        colWidths=[36],
        rowHeights=[36]
    )
    logo_inner.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, -1), colors.HexColor('#7F77DD')),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING',    (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))

    logo_text = Table(
        [[Paragraph('DocForge Hub', ParagraphStyle(
            'LogoName',
            parent=styles['Normal'],
            fontSize=13,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor('#1a1a3a'),
            leading=16
        ))],
        [Paragraph('AI Document Generation System', ParagraphStyle(
            'LogoSub',
            parent=styles['Normal'],
            fontSize=9,
            fontName='Helvetica',
            textColor=colors.HexColor('#888888'),
            leading=13
        ))]],
        colWidths=[220]
    )
    logo_text.setStyle(TableStyle([
        ('TOPPADDING',    (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('LEFTPADDING',   (0, 0), (-1, -1), 10),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
    ]))

    logo_table = Table(
        [[logo_inner, logo_text]],
        colWidths=[40, 220],
        hAlign='LEFT'
    )
    logo_table.setStyle(TableStyle([
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ('TOPPADDING',    (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(logo_table)
    story.append(Spacer(1, 80))

    story.append(Paragraph(
        doc_type.upper() if doc_type else "DOCUMENT",
        cover_doctype_style
    ))
    story.append(Spacer(1, 8))

    story.append(Paragraph(html.escape(title), cover_title_style))
    story.append(Spacer(1, 12))

    divider = Table(
        [['']],
        colWidths=[60],
        rowHeights=[3],
        hAlign='CENTER'
    )
    divider.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, -1), colors.HexColor('#7F77DD')),
        ('TOPPADDING',    (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(divider)
    story.append(Spacer(1, 20))

    if company_name:
        story.append(Paragraph(html.escape(company_name), cover_company_style))
        story.append(Spacer(1, 8))

    story.append(Paragraph(
        f"This document establishes the official procedures and guidelines for "
        f"{html.escape(title.lower())} within the organization.",
        cover_desc_style
    ))
    story.append(Spacer(1, 40))

    meta_items = [
        ("DEPARTMENT",  department   or "—"),
        ("VERSION",     version      or "v1.0"),
        ("CREATED ON",  date_str     or "—"),
        ("INDUSTRY",    industry     or "—"),
    ]

    meta_cells = []
    for label, value in meta_items:
        cell = Table(
            [
                [Paragraph(label, meta_label_style)],
                [Paragraph(html.escape(str(value)), meta_value_style)]
            ],
            colWidths=[160]
        )
        cell.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), colors.HexColor('#f8f8ff')),
            ('LEFTPADDING',   (0, 0), (-1, -1), 12),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 12),
            ('TOPPADDING',    (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('BOX',           (0, 0), (-1, -1), 0.5, colors.HexColor('#e8e8f8')),
        ]))
        meta_cells.append(cell)

    meta_table = Table(
        [[meta_cells[0], meta_cells[1]],
         [meta_cells[2], meta_cells[3]]],
        colWidths=[180, 180],
        hAlign='CENTER'
    )
    meta_table.setStyle(TableStyle([
        ('LEFTPADDING',   (0, 0), (-1, -1), 6),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 6),
        ('TOPPADDING',    (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 80))

    story.append(HRFlowable(
        width="100%", thickness=0.5,
        color=colors.HexColor('#e8e8f8'),
        spaceAfter=10
    ))

    footer_table = Table(
        [[
            Paragraph('Generated by DocForge Hub', ParagraphStyle(
                'FL', parent=styles['Normal'],
                fontSize=9, fontName='Helvetica',
                textColor=colors.HexColor('#aaaaaa')
            )),
            Paragraph('CONFIDENTIAL', ParagraphStyle(
                'FC', parent=styles['Normal'],
                fontSize=9, fontName='Helvetica-Bold',
                textColor=colors.HexColor('#7F77DD'),
                alignment=TA_CENTER
            )),
            Paragraph('Page 1', ParagraphStyle(
                'FR', parent=styles['Normal'],
                fontSize=9, fontName='Helvetica',
                textColor=colors.HexColor('#aaaaaa'),
                alignment=2
            ))
        ]],
        colWidths=[doc.width/3, doc.width/3, doc.width/3]
    )
    footer_table.setStyle(TableStyle([
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ('TOPPADDING',    (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(footer_table)
    story.append(PageBreak())


    # DOCUMENT CONTENT

    def build_pdf_table(rows):
        if not rows:
            return None
        pdf_rows = []
        for i, row in enumerate(rows):
            pdf_row = []
            for cell in row:
                style = table_header_style if i == 0 else table_cell_style
                pdf_row.append(Paragraph(html.escape(str(cell)), style))
            pdf_rows.append(pdf_row)

        col_count = max(len(r) for r in pdf_rows)
        col_width  = doc.width / col_count

        t = Table(pdf_rows, colWidths=[col_width] * col_count)
        t.setStyle(TableStyle([
            ('BACKGROUND',     (0, 0), (-1, 0),  colors.HexColor('#7F77DD')),
            ('TEXTCOLOR',      (0, 0), (-1, 0),  colors.white),
            ('FONTNAME',       (0, 0), (-1, 0),  'Helvetica-Bold'),
            ('FONTSIZE',       (0, 0), (-1, 0),  10),
            ('BACKGROUND',     (0, 1), (-1, -1), colors.HexColor('#f5f5ff')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1),
             [colors.HexColor('#f5f5ff'), colors.white]),
            ('GRID',           (0, 0), (-1, -1), 0.5, colors.HexColor('#c0c0d8')),
            ('LINEBELOW',      (0, 0), (-1, 0),  1.5, colors.HexColor('#534AB7')),
            ('TOPPADDING',     (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING',  (0, 0), (-1, -1), 8),
            ('LEFTPADDING',    (0, 0), (-1, -1), 10),
            ('RIGHTPADDING',   (0, 0), (-1, -1), 10),
            ('VALIGN',         (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        return t

    for row in sections:
        sec_title   = row[0] or "Untitled Section"
        sec_content = row[1] or "No content available"

        story.append(CondPageBreak(inch * 2))
        story.append(Paragraph(html.escape(sec_title), section_style))
        story.append(HRFlowable(
            width="100%", thickness=0.5,
            color=colors.HexColor('#e0e0f0'),
            spaceAfter=8
        ))

        lines     = sec_content.split('\n')
        i         = 0
        table_buf = []

        while i < len(lines):
            line = lines[i]

            if is_table_row(line):
                table_buf.append(line)
                i += 1
                while i < len(lines) and (
                    is_table_row(lines[i]) or is_separator_row(lines[i])
                ):
                    table_buf.append(lines[i])
                    i += 1
                table_rows = parse_table_rows(table_buf)
                if table_rows:
                    pdf_table = build_pdf_table(table_rows)
                    if pdf_table:
                        story.append(Spacer(1, 8))
                        story.append(pdf_table)
                        story.append(Spacer(1, 12))
                table_buf = []
                continue

            clean_line = line.strip()

            if not clean_line:
                story.append(Spacer(1, 4))
                i += 1
                continue

            clean_line = re.sub(r'#{1,6}\s*', '', clean_line)

            if clean_line.startswith('•') or \
               (clean_line.startswith('*') and not clean_line.startswith('**')) or \
               re.match(r'^-\s+[A-Za-z]', clean_line):
                bullet_text = re.sub(r'^[•*\-]\s*', '', clean_line)
                bullet_text = re.sub(
                    r'^([A-Za-z][A-Za-z0-9\s\(\)\/\-\&]{2,120}):\s',
                    r'**\1:** ',
                    bullet_text
                )
                rich = render_rich_text_pdf(bullet_text)
                story.append(Paragraph(f"• {rich}", bullet_style))
            else:
                rich = render_rich_text_pdf(clean_line)
                story.append(Paragraph(rich, body_style))

            i += 1

        story.append(Spacer(1, 12))

    # Page number function 
    def add_page_number(canvas, doc):
        page_num = canvas.getPageNumber()
        if page_num == 1:
            return
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        canvas.setFillColor(colors.HexColor('#aaaaaa'))
        canvas.drawString(60, 30, "DocForge Hub")
        canvas.drawCentredString(
            doc.pagesize[0] / 2, 30,
            f"Page {page_num}"
        )
        canvas.setFont('Helvetica-Bold', 9)
        canvas.setFillColor(colors.HexColor('#7F77DD'))
        canvas.drawRightString(doc.pagesize[0] - 60, 30, "CONFIDENTIAL")
        canvas.setStrokeColor(colors.HexColor('#7F77DD'))
        canvas.setLineWidth(2)
        canvas.line(
            60,
            doc.pagesize[1] - 40,
            doc.pagesize[0] - 60,
            doc.pagesize[1] - 40
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)

    return FileResponse(
        path=file_path,
        filename=file_name,
        media_type="application/pdf"
    )


# DOCX DOWNLOAD

@router.get("/download/docx/{document_id}")
def download_docx(document_id: str):
    conn   = get_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT dt.name FROM documents d
        JOIN document_templates dt ON d.template_id = dt.id
        WHERE d.id=%s
        """,
        (document_id,)
    )
    result = cursor.fetchone()
    if not result:
        raise HTTPException(status_code=404, detail="Document not found")
    title = result[0]

    cursor.execute(
        """
        SELECT DISTINCT ON (section_order)
            section_title, section_content, section_order
        FROM document_sections
        WHERE document_id=%s
        ORDER BY section_order, id DESC
        """,
        (document_id,)
    )
    sections = cursor.fetchall()
    cursor.close()
    conn.close()

    if not sections:
        raise HTTPException(status_code=404, detail="No content found")

    conn2   = get_connection()
    cursor2 = conn2.cursor()
    cursor2.execute(
        """
        SELECT
            dep.name,
            dty.name,
            cc.company_name,
            dt.industry,
            d.version,
            d.created_at
        FROM documents d
        JOIN document_templates dt ON d.template_id = dt.id
        JOIN departments dep ON dt.department_id = dep.id
        JOIN document_types dty ON dt.document_type_id = dty.id
        LEFT JOIN company_context cc ON d.company_id = cc.id
        WHERE d.id = %s
        """,
        (document_id,)
    )
    meta_row = cursor2.fetchone()
    cursor2.close()
    conn2.close()

    department   = meta_row[0] if meta_row else ""
    doc_type     = meta_row[1] if meta_row else "Document"
    company_name = meta_row[2] if meta_row else ""
    industry     = meta_row[3] if meta_row else ""
    version      = meta_row[4] if meta_row else "v1.0"
    created_at   = meta_row[5] if meta_row else datetime.now()
    date_str     = created_at.strftime("%B %d, %Y") if created_at else ""

    doc = Document()

    # from docx.shared import Pt, RGBColor, Cm
    # from docx.enum.text import WD_ALIGN_PARAGRAPH
    # from docx.oxml.ns import qn
    # from docx.oxml import OxmlElement

    # Cover page 
    cover_bar     = doc.add_paragraph()
    cover_bar.add_run("  ")
    cover_bar.paragraph_format.space_before = Pt(0)
    cover_bar.paragraph_format.space_after  = Pt(0)
    cover_bar_fmt = cover_bar._p.get_or_add_pPr()
    cover_bar_shd = OxmlElement('w:shd')
    cover_bar_shd.set(qn('w:val'),   'clear')
    cover_bar_shd.set(qn('w:color'), 'auto')
    cover_bar_shd.set(qn('w:fill'),  '7F77DD')
    cover_bar_fmt.append(cover_bar_shd)

    doc.add_paragraph("")
    doc.add_paragraph("")

    logo_para           = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run            = logo_para.add_run("D  DocForge Hub")
    logo_run.bold       = True
    logo_run.font.size  = Pt(16)
    logo_run.font.color.rgb = RGBColor(0x7F, 0x77, 0xDD)

    sub_para           = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_run            = sub_para.add_run("AI Document Generation System")
    sub_run.font.size  = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    dtype_para           = doc.add_paragraph()
    dtype_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dtype_run            = dtype_para.add_run(
        doc_type.upper() if doc_type else "DOCUMENT"
    )
    dtype_run.bold       = True
    dtype_run.font.size  = Pt(10)
    dtype_run.font.color.rgb = RGBColor(0x7F, 0x77, 0xDD)

    doc.add_paragraph("")

    title_para           = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run            = title_para.add_run(title)
    title_run.bold       = True
    title_run.font.size  = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x3a)

    doc.add_paragraph("")

    if company_name:
        comp_para           = doc.add_paragraph()
        comp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        comp_run            = comp_para.add_run(company_name)
        comp_run.font.size  = Pt(15)
        comp_run.font.color.rgb = RGBColor(0x4a, 0x4a, 0x6a)

    doc.add_paragraph("")

    desc_para           = doc.add_paragraph()
    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run            = desc_para.add_run(
        f"This document establishes the official procedures and guidelines "
        f"for {title.lower()} within the organization."
    )
    desc_run.font.size  = Pt(11)
    desc_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")
    doc.add_paragraph("")

    meta_tbl       = doc.add_table(rows=2, cols=2)
    meta_tbl.style = 'Table Grid'
    meta_data      = [
        ("DEPARTMENT", department or "—"),
        ("VERSION",    version    or "v1.0"),
        ("CREATED ON", date_str   or "—"),
        ("INDUSTRY",   industry   or "—"),
    ]
    for (r, c), (label, value) in zip(
        [(0,0),(0,1),(1,0),(1,1)], meta_data
    ):
        cell = meta_tbl.cell(r, c)
        cell.paragraphs[0].clear()
        lbl_run      = cell.paragraphs[0].add_run(label)
        lbl_run.bold = True
        lbl_run.font.size      = Pt(8)
        lbl_run.font.color.rgb = RGBColor(0x7F, 0x77, 0xDD)
        val_para = cell.add_paragraph()
        val_run  = val_para.add_run(value)
        val_run.bold       = True
        val_run.font.size  = Pt(11)
        val_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x3a)
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F8F8FF')
        tcPr.append(shd)

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    footer_para           = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run            = footer_para.add_run(
        "Generated by DocForge Hub  |  CONFIDENTIAL"
    )
    footer_run.font.size  = Pt(9)
    footer_run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    doc.add_page_break()

    # Document content
    def add_rich_line_docx(para, text: str):
        parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run      = para.add_run(part[2:-2])
                run.bold = True
                run.font.size      = Pt(11)
                run.font.color.rgb = RGBColor(0x2a, 0x2a, 0x4a)
            elif part.startswith('__') and part.endswith('__'):
                run           = para.add_run(part[2:-2])
                run.underline = True
                run.font.size      = Pt(11)
                run.font.color.rgb = RGBColor(0x2a, 0x2a, 0x4a)
            else:
                if part:
                    run = para.add_run(part)
                    run.font.size      = Pt(11)
                    run.font.color.rgb = RGBColor(0x2a, 0x2a, 0x4a)

    def add_table_to_doc(doc, rows):
        if not rows:
            return
        col_count  = max(len(r) for r in rows)
        normalized = []
        for row in rows:
            while len(row) < col_count:
                row.append('')
            normalized.append(row[:col_count])

        table       = doc.add_table(rows=len(normalized), cols=col_count)
        table.style = 'Table Grid'

        for i, row_data in enumerate(normalized):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell      = row.cells[j]
                cell.text = cell_text
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold      = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size      = Pt(10)
                    tc   = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd  = OxmlElement('w:shd')
                    shd.set(qn('w:val'),   'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'),  '7F77DD')
                    tcPr.append(shd)
                else:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size      = Pt(10)
                            run.font.color.rgb = RGBColor(0x2a, 0x2a, 0x4a)
                    if i % 2 == 0:
                        tc   = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd  = OxmlElement('w:shd')
                        shd.set(qn('w:val'),   'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'),  'F5F5FF')
                        tcPr.append(shd)
        doc.add_paragraph("")

    for row in sections:
        sec_title   = row[0] or "Untitled Section"
        sec_content = row[1] or "No content available"

        heading = doc.add_heading(sec_title, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x7F, 0x77, 0xDD)
            run.font.size      = Pt(14)

        lines     = sec_content.split('\n')
        i         = 0
        table_buf = []

        while i < len(lines):
            line = lines[i]

            if is_table_row(line):
                table_buf.append(line)
                i += 1
                while i < len(lines) and (
                    is_table_row(lines[i]) or is_separator_row(lines[i])
                ):
                    table_buf.append(lines[i])
                    i += 1
                table_rows = parse_table_rows(table_buf)
                if table_rows:
                    add_table_to_doc(doc, table_rows)
                table_buf = []
                continue

            clean_line = line.strip()

            if not clean_line:
                doc.add_paragraph("")
                i += 1
                continue

            clean_line = re.sub(r'#{1,6}\s*', '', clean_line).strip()

            if clean_line.startswith('•') or \
               (clean_line.startswith('*') and not clean_line.startswith('**')) or \
               re.match(r'^-\s+[A-Za-z]', clean_line):
                bullet_text = re.sub(r'^[•*\-]\s*', '', clean_line)
                bullet_text = re.sub(
                    r'^([A-Za-z][A-Za-z0-9\s\(\)\/\-\&]{2,120}):\s',
                    r'**\1:** ',
                    bullet_text
                )
                para        = doc.add_paragraph(style='Normal')
                para.paragraph_format.left_indent = Pt(16)
                run_bullet           = para.add_run('• ')
                run_bullet.font.size = Pt(11)
                add_rich_line_docx(para, bullet_text)
            else:
                clean_line = re.sub(
                    r'^([A-Za-z][A-Za-z0-9\s\(\)\/\-\&]{2,120}):\s',
                    r'**\1:** ',
                    clean_line
                )
                para = doc.add_paragraph(style='Normal')
                add_rich_line_docx(para, clean_line)

            i += 1

        doc.add_paragraph("")

    file_name = title.lower().replace(" ", "_") + ".docx"
    file_path = f"/tmp/{file_name}"
    doc.save(file_path)

    return FileResponse(
        path=file_path,
        filename=file_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )